"""Positive / negative / ambiguous detection cases + auth flow."""

import io
import uuid

import pymupdf
from docx import Document

from app.schemas import DetectedEntity, EntityLocation


def entity_types(resp):
    return {e["entity_type"] for e in resp.json()["detected_entities"]}


# ---------------------------------------------------------------- auth ----


def test_register_returns_api_key(client):
    resp = client.post("/register", json={"email": f"x-{uuid.uuid4().hex[:8]}@sensen.dev"})
    assert resp.status_code == 200
    assert "api_key" in resp.json()


def test_duplicate_email_rejected(client):
    email = f"dup-{uuid.uuid4().hex[:8]}@sensen.dev"
    assert client.post("/register", json={"email": email}).status_code == 200
    assert client.post("/register", json={"email": email}).status_code == 400


def test_scan_without_key_rejected(client):
    resp = client.post("/api/v1/scan", json={"text": "hello"})
    assert resp.status_code == 422  # missing required header


def test_scan_with_invalid_key_rejected(client):
    resp = client.post(
        "/api/v1/scan", json={"text": "hello"}, headers={"X-API-Key": "not-a-real-key"}
    )
    assert resp.status_code == 401


def test_unsupported_language_rejected(scan):
    resp = scan("hello", language="vi")
    assert resp.status_code == 400


# ------------------------------------------------------------ positive ----


def test_detects_email(scan):
    resp = scan("Reach me at john.doe@example.com please.")
    assert resp.status_code == 200
    assert "EMAIL_ADDRESS" in entity_types(resp)


def test_detects_contract_id_with_context(scan):
    resp = scan("Hợp đồng số HD-2026-0142 đã được ký.", confidence_threshold=0.5)
    assert "CONTRACT_ID" in entity_types(resp)


def test_detects_aws_access_key(scan):
    resp = scan("Leaked credential: AKIAABCDEFGHIJKLMNOP was found in the log file.")
    assert "INFRA_SECRET" in entity_types(resp)


def test_detects_db_connection_string(scan):
    resp = scan("conn_str = postgresql://admin:pass123@db.internal:5432/prod")
    assert "INFRA_SECRET" in entity_types(resp)


def test_detects_financial_metric_with_context(scan):
    resp = scan(
        "Lương thưởng tháng này là 50.000.000 VND cho toàn bộ nhân viên.",
        confidence_threshold=0.5,
    )
    assert "FINANCIAL_METRIC" in entity_types(resp)


def test_detects_ip_confidentiality_marker(scan):
    resp = scan(
        "Tài liệu này CONFIDENTIAL, chứa source code độc quyền của công ty.",
        confidence_threshold=0.3,
    )
    assert "IP_SENSITIVE_MARKER" in entity_types(resp)


def test_anonymize_masks_email(scan):
    resp = scan("Contact test@example.com now", anonymize=True)
    body = resp.json()
    assert "test@example.com" not in body["anonymized_content"]["text"]
    assert "<EMAIL_ADDRESS>" in body["anonymized_content"]["text"]


# ------------------------------------------------------------ negative ----


def test_generic_sentence_has_no_high_confidence_hits(scan):
    # Avoids names/dates/orgs so spaCy's NER has nothing to misfire on.
    resp = scan("The quick brown fox jumps over the lazy dog.", confidence_threshold=0.7)
    assert resp.json()["detected_entities"] == []


def test_short_numbers_not_flagged_as_contract(scan):
    # A random number sequence must not be mistaken for a sensitive code.
    resp = scan("Row values: 42, 17, 93, 8", confidence_threshold=0.7)
    assert "CONTRACT_ID" not in entity_types(resp)


# ----------------------------------------------------------- ambiguous ----


def test_generic_code_shape_filtered_at_default_threshold_without_context(scan):
    # Same shape as a contract id but with no legal context nearby.
    text = "Reference AB-12-XY9Z used for internal tracking only."
    strict = scan(text, confidence_threshold=0.7)
    loose = scan(text, confidence_threshold=0.2)
    assert "CONTRACT_ID" not in entity_types(strict)
    assert "CONTRACT_ID" in entity_types(loose)


def test_mobile_number_not_misread_as_tax_code(scan):
    # VN mobile numbers must not collide with INTERNAL_TAX_CODE (see recognizers.yaml).
    resp = scan("Gọi tôi qua số 0912345678 nhé.", confidence_threshold=0.2)
    assert "INTERNAL_TAX_CODE" not in entity_types(resp)


# --------------------------------------------------------- file upload ----


def test_scan_pdf_file_extracts_and_detects(client, api_key):
    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Contact john@example.com for details.")
    raw = doc.tobytes()
    doc.close()

    resp = client.post(
        "/api/v1/scan/file",
        files={"file": ("contract.pdf", raw, "application/pdf")},
        headers={"X-API-Key": api_key},
    )
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["document_metadata"]["file_type"] == "pdf"
    assert "EMAIL_ADDRESS" in entity_types(resp)


def test_scan_docx_file_extracts_and_detects(client, api_key):
    document = Document()
    document.add_paragraph("Hợp đồng số HD-2026-0142 đã được ký kết.")
    buf = io.BytesIO()
    document.save(buf)

    resp = client.post(
        "/api/v1/scan/file",
        files={
            "file": (
                "contract.docx",
                buf.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"confidence_threshold": "0.5"},
        headers={"X-API-Key": api_key},
    )
    assert resp.status_code == 200, resp.text
    assert "CONTRACT_ID" in entity_types(resp)


def test_scan_unsupported_file_type_rejected(client, api_key):
    resp = client.post(
        "/api/v1/scan/file",
        files={"file": ("notes.csv", b"a,b,c", "text/csv")},
        headers={"X-API-Key": api_key},
    )
    assert resp.status_code == 422


def test_corrupt_pdf_rejected_cleanly_not_a_500(client, api_key):
    # pymupdf.FileDataError must map to a 422, not an unhandled 500.
    resp = client.post(
        "/api/v1/scan/file",
        files={"file": ("garbage.pdf", b"this is not a valid pdf at all", "application/pdf")},
        headers={"X-API-Key": api_key},
    )
    assert resp.status_code == 422
    assert "pdf" in resp.json()["detail"].lower()


def test_empty_pdf_rejected_cleanly_not_a_500(client, api_key):
    resp = client.post(
        "/api/v1/scan/file",
        files={"file": ("empty.pdf", b"", "application/pdf")},
        headers={"X-API-Key": api_key},
    )
    assert resp.status_code == 422


def test_corrupt_docx_rejected_cleanly_not_a_500(client, api_key):
    # Document() raises zipfile.BadZipFile or KeyError depending on how
    # malformed the file is -- both must map to a 422.
    resp = client.post(
        "/api/v1/scan/file",
        files={
            "file": (
                "garbage.docx",
                b"not a real docx zip file",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        headers={"X-API-Key": api_key},
    )
    assert resp.status_code == 422
    assert "docx" in resp.json()["detail"].lower()


def _build_image_only_pdf(num_pages: int = 1) -> bytes:
    # A page with only an inserted image, no text layer -- same shape as a
    # real scanned document, exercises app/extract.py's OCR fallback path.
    src = pymupdf.open()
    src.new_page(width=200, height=100)
    pix = src[0].get_pixmap()
    img_bytes = pix.tobytes("png")
    src.close()

    out = pymupdf.open()
    for _ in range(num_pages):
        page = out.new_page(width=pix.width, height=pix.height)
        page.insert_image(page.rect, stream=img_bytes)
    raw = out.tobytes()
    out.close()
    return raw


def test_scanned_pdf_falls_back_to_ocr(client, api_key, monkeypatch):
    monkeypatch.setattr(
        "app.extract.pytesseract.image_to_string",
        lambda img, lang=None: "Hợp đồng số HD-2026-7777",
    )
    resp = client.post(
        "/api/v1/scan/file",
        files={"file": ("scanned.pdf", _build_image_only_pdf(), "application/pdf")},
        headers={"X-API-Key": api_key},
    )
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["document_metadata"]["processing_mode"] == "ocr_local"
    assert "CONTRACT_ID" in entity_types(resp)


def test_scanned_pdf_without_tesseract_installed_gives_clear_422(client, api_key, monkeypatch):
    import pytesseract

    def _raise_not_found(img, lang=None):
        raise pytesseract.TesseractNotFoundError()

    monkeypatch.setattr("app.extract.pytesseract.image_to_string", _raise_not_found)
    resp = client.post(
        "/api/v1/scan/file",
        files={"file": ("scanned.pdf", _build_image_only_pdf(), "application/pdf")},
        headers={"X-API-Key": api_key},
    )
    assert resp.status_code == 422
    assert "tesseract" in resp.json()["detail"].lower()


def test_scanned_pdf_over_page_limit_rejected_before_ocr_attempt(client, api_key, monkeypatch):
    from app.extract import MAX_OCR_PAGES

    def _boom(img, lang=None):
        raise AssertionError("OCR must not run once the page-count cap is exceeded")

    monkeypatch.setattr("app.extract.pytesseract.image_to_string", _boom)
    resp = client.post(
        "/api/v1/scan/file",
        files={
            "file": (
                "scanned.pdf",
                _build_image_only_pdf(num_pages=MAX_OCR_PAGES + 1),
                "application/pdf",
            )
        },
        headers={"X-API-Key": api_key},
    )
    assert resp.status_code == 422
    assert str(MAX_OCR_PAGES) in resp.json()["detail"]


# --------------------------------------------------------- PDF redaction ----
# /api/v1/redact/file returns an actual redacted file, not just masked text.


def _build_scanned_pdf_with_real_text(text: str) -> bytes:
    src = pymupdf.open()
    p = src.new_page()
    p.insert_text((72, 72), text)
    pix = p.get_pixmap(dpi=200)
    img_bytes = pix.tobytes("png")
    src.close()

    out = pymupdf.open()
    page = out.new_page(width=pix.width, height=pix.height)
    page.insert_image(page.rect, stream=img_bytes)
    raw = out.tobytes()
    out.close()
    return raw


def test_redact_digital_pdf_removes_sensitive_text_for_real(client, api_key):
    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Contact john@example.com or call 0912345678 for details.")
    raw = doc.tobytes()
    doc.close()

    resp = client.post(
        "/api/v1/redact/file",
        files={"file": ("contract.pdf", raw, "application/pdf")},
        data={"confidence_threshold": "0.3"},
        headers={"X-API-Key": api_key},
    )
    assert resp.status_code == 200, resp.text
    assert resp.headers["content-type"] == "application/pdf"
    assert "attachment" in resp.headers["content-disposition"]

    redacted = pymupdf.open(stream=resp.content, filetype="pdf")
    text = redacted[0].get_text()
    assert "john@example.com" not in text
    assert "0912345678" not in text


def test_redact_scanned_pdf_local_blacks_out_the_image(client, api_key):
    import pytesseract
    from PIL import Image

    raw = _build_scanned_pdf_with_real_text(
        "Contact john@example.com or call 0912345678 for details."
    )
    resp = client.post(
        "/api/v1/redact/file",
        files={"file": ("scanned.pdf", raw, "application/pdf")},
        data={"confidence_threshold": "0.3", "ocr_engine": "local"},
        headers={"X-API-Key": api_key},
    )
    assert resp.status_code == 200, resp.text

    redacted = pymupdf.open(stream=resp.content, filetype="pdf")
    pix = redacted[0].get_pixmap(dpi=200)
    mode = "RGBA" if pix.alpha else "RGB"
    img = Image.frombytes(mode, (pix.width, pix.height), pix.samples)
    retext = pytesseract.image_to_string(img, lang="eng")
    assert "john@example.com" not in retext
    assert "0912345678" not in retext


def test_redact_output_is_reasonably_sized_not_bloated(client, api_key):
    # doc.tobytes() without garbage/deflate/clean left old page data behind.
    raw = _build_scanned_pdf_with_real_text("Contact john@example.com for details.")
    resp = client.post(
        "/api/v1/redact/file",
        files={"file": ("scanned.pdf", raw, "application/pdf")},
        data={"confidence_threshold": "0.3", "ocr_engine": "local"},
        headers={"X-API-Key": api_key},
    )
    assert resp.status_code == 200, resp.text
    assert len(resp.content) < 500_000, f"redacted PDF suspiciously large: {len(resp.content)} bytes"


def test_redact_rejects_unsupported_file_type(client, api_key):
    resp = client.post(
        "/api/v1/redact/file",
        files={"file": ("notes.csv", b"a,b,c", "text/csv")},
        headers={"X-API-Key": api_key},
    )
    assert resp.status_code == 422
    assert ".csv" in resp.json()["detail"] or "csv" in resp.json()["detail"].lower()


def test_redact_txt_masks_sensitive_text(client, api_key):
    resp = client.post(
        "/api/v1/redact/file",
        files={
            "file": (
                "notes.txt",
                b"john@example.com is the contact email.",
                "text/plain",
            )
        },
        data={"confidence_threshold": "0.3"},
        headers={"X-API-Key": api_key},
    )
    assert resp.status_code == 200, resp.text
    assert resp.headers["content-type"].startswith("text/plain")
    body = resp.content.decode("utf-8")
    assert "john@example.com" not in body
    assert "<EMAIL_ADDRESS>" in body


def test_redact_docx_removes_sensitive_text_for_real(client, api_key):
    import zipfile

    buf = io.BytesIO()
    document = Document()
    document.add_paragraph("john@example.com is the contact email.")
    document.save(buf)

    resp = client.post(
        "/api/v1/redact/file",
        files={
            "file": (
                "contract.docx",
                buf.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"confidence_threshold": "0.3"},
        headers={"X-API-Key": api_key},
    )
    assert resp.status_code == 200, resp.text
    assert (
        resp.headers["content-type"]
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    redacted = Document(io.BytesIO(resp.content))
    assert "john@example.com" not in redacted.paragraphs[0].text

    # Real XML-level deletion, not just the friendly API's view.
    z = zipfile.ZipFile(io.BytesIO(resp.content))
    xml = z.read("word/document.xml").decode("utf-8")
    assert "john@example.com" not in xml


def test_redact_docx_covers_tables_and_headers(client, api_key):
    buf = io.BytesIO()
    document = Document()
    document.add_paragraph("Body text with nothing sensitive.")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "jane.doe@example.com is in this cell."
    document.sections[0].header.paragraphs[0].text = "Memo phone: 0987654321."
    document.save(buf)

    resp = client.post(
        "/api/v1/redact/file",
        files={
            "file": (
                "memo.docx",
                buf.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"confidence_threshold": "0.3"},
        headers={"X-API-Key": api_key},
    )
    assert resp.status_code == 200, resp.text
    redacted = Document(io.BytesIO(resp.content))
    assert "jane.doe@example.com" not in redacted.tables[0].cell(0, 0).text
    assert "0987654321" not in redacted.sections[0].header.paragraphs[0].text


def test_redact_docx_embedded_image_blacks_out_the_image(client, api_key):
    import pymupdf
    import pytesseract
    from PIL import Image

    src = pymupdf.open()
    page = src.new_page()
    page.insert_text((36, 36), "Contact john@example.com or call 0912345678.", fontsize=14)
    pix = page.get_pixmap(dpi=200)
    img_bytes = pix.tobytes("png")
    src.close()

    document = Document()
    document.add_paragraph("Nothing sensitive in the body.")
    document.add_picture(io.BytesIO(img_bytes))
    buf = io.BytesIO()
    document.save(buf)

    resp = client.post(
        "/api/v1/redact/file",
        files={
            "file": (
                "scanned.docx",
                buf.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"confidence_threshold": "0.3", "ocr_engine": "local"},
        headers={"X-API-Key": api_key},
    )
    assert resp.status_code == 200, resp.text

    redacted = Document(io.BytesIO(resp.content))
    image_part = next(
        p for p in redacted.part.related_parts.values() if p.content_type.startswith("image/")
    )
    img2 = Image.open(io.BytesIO(image_part.blob))
    retext = pytesseract.image_to_string(img2, lang="eng")
    assert "john@example.com" not in retext
    assert "0912345678" not in retext


def test_redact_docx_unsupported_embedded_image_format_gives_clear_422(client, api_key, monkeypatch):
    from app import redact as redact_module

    monkeypatch.delitem(redact_module._CONTENT_TYPE_TO_PIL_FORMAT, "image/png")

    document = Document()
    document.add_paragraph("Nothing sensitive in the body.")
    # A 1x1 PNG is enough -- this test is about the content-type gate,
    # not what's actually drawn in the image.
    from PIL import Image as PILImage

    tiny = io.BytesIO()
    PILImage.new("RGB", (1, 1)).save(tiny, format="PNG")
    tiny.seek(0)
    document.add_picture(tiny)
    buf = io.BytesIO()
    document.save(buf)

    resp = client.post(
        "/api/v1/redact/file",
        files={
            "file": (
                "scanned.docx",
                buf.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"confidence_threshold": "0.3"},
        headers={"X-API-Key": api_key},
    )
    assert resp.status_code == 422
    assert "image/png" in resp.json()["detail"]


def test_redact_via_cloud_engine_consumes_ocr_api_quota(client, monkeypatch):
    from PIL import Image

    from app import redact as redact_module
    from app.ocr_api import OcrWord

    mocked_bbox = (100.0, 100.0, 400.0, 130.0)
    monkeypatch.setattr(
        redact_module,
        "ocr_words_via_api",
        lambda engine, png_bytes, model=None: [
            OcrWord(
                text="Contact john@example.com or call 0912345678 for details.",
                bbox=mocked_bbox,
            )
        ],
    )
    email = f"redact-quota-{uuid.uuid4().hex[:8]}@sensen.dev"
    key = client.post("/register", json={"email": email}).json()["api_key"]
    raw = _build_scanned_pdf_with_real_text("placeholder -- OCR is mocked")

    resp = client.post(
        "/api/v1/redact/file",
        files={"file": ("scanned.pdf", raw, "application/pdf")},
        data={"confidence_threshold": "0.3", "ocr_engine": "gemini"},
        headers={"X-API-Key": key},
    )
    assert resp.status_code == 200, resp.text

    # get_text() is always "" for an image page -- check the actual pixel
    # color at the mocked bbox instead.
    redacted = pymupdf.open(stream=resp.content, filetype="pdf")
    pix = redacted[0].get_pixmap(dpi=200)
    mode = "RGBA" if pix.alpha else "RGB"
    img = Image.frombytes(mode, (pix.width, pix.height), pix.samples)
    cx, cy = int((mocked_bbox[0] + mocked_bbox[2]) / 2), int((mocked_bbox[1] + mocked_bbox[3]) / 2)
    pixel = img.convert("RGB").getpixel((cx, cy))
    assert pixel == (0, 0, 0), f"expected black at redacted bbox center, got {pixel}"

    usage = client.get("/api/v1/usage", headers={"X-API-Key": key}).json()
    assert usage["ocr_api_used"] == 1


def test_redact_cloud_engine_no_boxes_returned_gives_clear_422(client, api_key, monkeypatch):
    from app import redact as redact_module

    monkeypatch.setattr(redact_module, "ocr_words_via_api", lambda engine, png_bytes, model=None: [])
    raw = _build_scanned_pdf_with_real_text("Contact john@example.com for details.")
    resp = client.post(
        "/api/v1/redact/file",
        files={"file": ("scanned.pdf", raw, "application/pdf")},
        data={"confidence_threshold": "0.3", "ocr_engine": "gemini"},
        headers={"X-API-Key": api_key},
    )
    assert resp.status_code == 422
    assert "no usable text/boxes" in resp.json()["detail"]


def test_redact_digital_page_fails_safely_when_entity_cant_be_located():
    # An unlocatable entity must fail loudly (RedactionFailed), never
    # silently produce a PDF that looks redacted but missed something.
    from app.engine import build_engines
    from app.redact import RedactionFailed, _redact_digital_page

    analyzer, _ = build_engines()

    class StubPage:
        number = 0

        def get_text(self):
            # No leading capitalized word -- avoids a known underthesea
            # false positive that would add an unrelated entity first.
            return "john@example.com is the contact email."

        def search_for(self, text):
            return []  # simulates pymupdf failing to locate the text on the page

        def add_redact_annot(self, *a, **k):
            raise AssertionError("must not attempt to redact when search_for found nothing")

        def apply_redactions(self):
            raise AssertionError("must not apply redactions when search_for found nothing")

    try:
        _redact_digital_page(StubPage(), analyzer, 0.3, False, None)
        assert False, "expected RedactionFailed when the entity can't be located on the page"
    except RedactionFailed as exc:
        assert "john@example.com" in str(exc)


# --------------------------------------------------------- cloud ocr api ----
# Patches target app.extract.ocr_image_via_api, the bound name it imports.


def test_cloud_ocr_engine_used_when_requested(client, api_key, monkeypatch):
    captured = {}

    def _fake_ocr(engine, png_bytes, model=None):
        captured["engine"] = engine
        captured["model"] = model
        return "Hợp đồng số HD-2026-7777"

    monkeypatch.setattr("app.extract.ocr_image_via_api", _fake_ocr)
    resp = client.post(
        "/api/v1/scan/file",
        files={"file": ("scanned.pdf", _build_image_only_pdf(), "application/pdf")},
        data={"ocr_engine": "gemini"},
        headers={"X-API-Key": api_key},
    )
    assert resp.status_code == 200, resp.text
    assert captured["engine"] == "gemini"
    assert captured["model"] is None  # no override passed -> falls back to the env-var default
    assert resp.json()["document_metadata"]["processing_mode"] == "ocr_gemini"
    assert "CONTRACT_ID" in entity_types(resp)


def test_cloud_ocr_model_override_passed_through(client, api_key, monkeypatch):
    # Deliberately a different model than deep_scan's own DEFAULT_MODEL_ID
    # -- OCR and deep_scan must be independently overridable per request.
    captured = {}

    def _fake_ocr(engine, png_bytes, model=None):
        captured["model"] = model
        return "Hợp đồng số HD-2026-7777"

    monkeypatch.setattr("app.extract.ocr_image_via_api", _fake_ocr)
    resp = client.post(
        "/api/v1/scan/file",
        files={"file": ("scanned.pdf", _build_image_only_pdf(), "application/pdf")},
        data={"ocr_engine": "gemini", "ocr_model": "gemini-flash-latest"},
        headers={"X-API-Key": api_key},
    )
    assert resp.status_code == 200, resp.text
    assert captured["model"] == "gemini-flash-latest"


def test_ocr_models_endpoint_reuses_deep_scan_listing_for_gemini(client, api_key, monkeypatch):
    monkeypatch.delenv("LANGEXTRACT_API_KEY", raising=False)
    resp = client.get(
        "/api/v1/ocr/models", params={"engine": "gemini"}, headers={"X-API-Key": api_key}
    )
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "skipped_no_key"
    assert body["default_model"] == "gemini-flash-lite-latest"


def test_ocr_models_endpoint_reports_no_key_for_openai(client, api_key, monkeypatch):
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)
    resp = client.get(
        "/api/v1/ocr/models", params={"engine": "openai"}, headers={"X-API-Key": api_key}
    )
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert body["status"] == "skipped_no_key"
    assert body["default_model"] == "gpt-5.6"
    assert body["models"] == []


def test_ocr_models_endpoint_rejects_unknown_engine(client, api_key):
    resp = client.get(
        "/api/v1/ocr/models", params={"engine": "deepseek"}, headers={"X-API-Key": api_key}
    )
    assert resp.status_code == 422


def test_ocr_models_endpoint_requires_auth(client):
    resp = client.get("/api/v1/ocr/models", params={"engine": "gemini"})
    assert resp.status_code == 422  # missing required header


def test_cloud_ocr_not_configured_gives_clear_422(client, api_key, monkeypatch):
    from app.ocr_api import OcrApiNotConfigured

    def _raise_not_configured(engine, png_bytes, model=None):
        raise OcrApiNotConfigured("openai", "OPENAI_API_KEY")

    monkeypatch.setattr("app.extract.ocr_image_via_api", _raise_not_configured)
    resp = client.post(
        "/api/v1/scan/file",
        files={"file": ("scanned.pdf", _build_image_only_pdf(), "application/pdf")},
        data={"ocr_engine": "openai"},
        headers={"X-API-Key": api_key},
    )
    assert resp.status_code == 422
    assert "OPENAI_API_KEY" in resp.json()["detail"]


def test_cloud_ocr_provider_failure_gives_clear_422(client, api_key, monkeypatch):
    from app.ocr_api import OcrApiError

    def _raise_error(engine, png_bytes, model=None):
        raise OcrApiError("simulated provider failure")

    monkeypatch.setattr("app.extract.ocr_image_via_api", _raise_error)
    resp = client.post(
        "/api/v1/scan/file",
        files={"file": ("scanned.pdf", _build_image_only_pdf(), "application/pdf")},
        data={"ocr_engine": "grok"},
        headers={"X-API-Key": api_key},
    )
    assert resp.status_code == 422
    assert "simulated provider failure" in resp.json()["detail"]


def test_invalid_ocr_engine_rejected(client, api_key):
    resp = client.post(
        "/api/v1/scan/file",
        files={"file": ("scanned.pdf", _build_image_only_pdf(), "application/pdf")},
        data={"ocr_engine": "deepseek"},
        headers={"X-API-Key": api_key},
    )
    assert resp.status_code == 422


def test_ocr_api_quota_exceeded_after_cap(client, monkeypatch):
    from app.pages import MAX_OCR_API_PER_KEY

    monkeypatch.setattr(
        "app.extract.ocr_image_via_api", lambda engine, png_bytes, model=None: "Hợp đồng số HD-2026-7777"
    )
    email = f"ocr-quota-{uuid.uuid4().hex[:8]}@sensen.dev"
    key = client.post("/register", json={"email": email}).json()["api_key"]

    for _ in range(MAX_OCR_API_PER_KEY):
        resp = client.post(
            "/api/v1/scan/file",
            files={"file": ("scanned.pdf", _build_image_only_pdf(), "application/pdf")},
            data={"ocr_engine": "gemini"},
            headers={"X-API-Key": key},
        )
        assert resp.status_code == 200, resp.text

    resp = client.post(
        "/api/v1/scan/file",
        files={"file": ("scanned.pdf", _build_image_only_pdf(), "application/pdf")},
        data={"ocr_engine": "gemini"},
        headers={"X-API-Key": key},
    )
    assert resp.status_code == 422
    assert "quota" in resp.json()["detail"].lower()


def test_ocr_api_quota_not_consumed_by_local_engine(client, monkeypatch):
    monkeypatch.setattr(
        "app.extract.pytesseract.image_to_string", lambda img, lang=None: "some text"
    )
    email = f"ocr-local-{uuid.uuid4().hex[:8]}@sensen.dev"
    key = client.post("/register", json={"email": email}).json()["api_key"]

    client.post(
        "/api/v1/scan/file",
        files={"file": ("scanned.pdf", _build_image_only_pdf(), "application/pdf")},
        headers={"X-API-Key": key},
    )
    resp = client.get("/api/v1/usage", headers={"X-API-Key": key})
    assert resp.json()["ocr_api_used"] == 0


def test_usage_endpoint_reports_counts(client, monkeypatch):
    monkeypatch.setattr("app.scanning.run_deep_scan", lambda text, model_id=None: ([], "ok"))
    monkeypatch.setattr(
        "app.extract.ocr_image_via_api", lambda engine, png_bytes, model=None: "Hợp đồng số HD-2026-7777"
    )
    email = f"usage-{uuid.uuid4().hex[:8]}@sensen.dev"
    key = client.post("/register", json={"email": email}).json()["api_key"]

    client.post(
        "/api/v1/scan", json={"text": "hello", "deep_scan": True}, headers={"X-API-Key": key}
    )
    client.post(
        "/api/v1/scan/file",
        files={"file": ("scanned.pdf", _build_image_only_pdf(), "application/pdf")},
        data={"ocr_engine": "openai"},
        headers={"X-API-Key": key},
    )
    resp = client.get("/api/v1/usage", headers={"X-API-Key": key})
    assert resp.status_code == 200
    body = resp.json()
    assert body["deep_scan_used"] == 1
    assert body["ocr_api_used"] == 1


def test_usage_endpoint_requires_auth(client):
    resp = client.get("/api/v1/usage")
    assert resp.status_code == 422  # missing required header


# ------------------------------------------- rate limiter (app/rate_limiter.py) ----
# time.monotonic/time.sleep are faked so tests run instantly.


def test_rate_limiter_allows_up_to_max_calls_without_blocking(monkeypatch):
    from app import rate_limiter

    fake_now = {"t": 0.0}
    monkeypatch.setattr(rate_limiter.time, "monotonic", lambda: fake_now["t"])
    sleeps = []
    monkeypatch.setattr(rate_limiter.time, "sleep", lambda s: sleeps.append(s))

    limiter = rate_limiter.SlidingWindowRateLimiter(max_calls=3, label="test")
    for _ in range(3):
        limiter.acquire()
    assert sleeps == []


def test_rate_limiter_blocks_when_budget_exhausted_then_recovers(monkeypatch):
    from app import rate_limiter

    fake_now = {"t": 0.0}

    def fake_sleep(s):
        fake_now["t"] += s  # simulate time actually passing while asleep

    monkeypatch.setattr(rate_limiter.time, "monotonic", lambda: fake_now["t"])
    monkeypatch.setattr(rate_limiter.time, "sleep", fake_sleep)

    limiter = rate_limiter.SlidingWindowRateLimiter(max_calls=2, label="test")
    limiter.acquire()
    limiter.acquire()
    limiter.acquire()  # budget exhausted, must wait for the window to free up
    assert fake_now["t"] > 0


def test_rate_limiter_disabled_when_max_calls_is_zero(monkeypatch):
    from app import rate_limiter

    sleeps = []
    monkeypatch.setattr(rate_limiter.time, "sleep", lambda s: sleeps.append(s))
    limiter = rate_limiter.SlidingWindowRateLimiter(max_calls=0, label="test")
    for _ in range(50):
        limiter.acquire()
    assert sleeps == []


def test_int_env_helper_parses_and_falls_back(monkeypatch):
    from app.rate_limiter import _int_env

    monkeypatch.setenv("SENSEN_TEST_RPM_VAR", "7")
    assert _int_env("SENSEN_TEST_RPM_VAR", 99) == 7
    monkeypatch.delenv("SENSEN_TEST_RPM_VAR", raising=False)
    assert _int_env("SENSEN_TEST_RPM_VAR", 99) == 99
    monkeypatch.setenv("SENSEN_TEST_RPM_VAR", "not-a-number")
    assert _int_env("SENSEN_TEST_RPM_VAR", 99) == 99


def test_ocr_api_and_deep_scan_share_the_same_gemini_rate_limiter():
    # Both draw down the same real Gemini RPM budget -- must share one instance.
    from app import deep_scan, ocr_api

    assert ocr_api.gemini_limiter is deep_scan.gemini_limiter


def test_ocr_gemini_call_goes_through_the_shared_rate_limiter(monkeypatch):
    from app import ocr_api

    calls = {"n": 0}
    monkeypatch.setattr(ocr_api.gemini_limiter, "acquire", lambda: calls.__setitem__("n", calls["n"] + 1))

    class FakeResponse:
        text = "extracted text"

    class FakeModels:
        def generate_content(self, **kwargs):
            return FakeResponse()

    class FakeClient:
        def __init__(self, **kwargs):
            self.models = FakeModels()

    monkeypatch.setattr(ocr_api.genai, "Client", FakeClient)
    monkeypatch.setenv("LANGEXTRACT_API_KEY", "fake-key")
    result = ocr_api.ocr_image_via_api("gemini", b"fake-png")
    assert result == "extracted text"
    assert calls["n"] == 1


def test_run_deep_scan_call_goes_through_the_shared_rate_limiter(monkeypatch):
    from app import deep_scan

    calls = {"n": 0}
    monkeypatch.setattr(deep_scan.gemini_limiter, "acquire", lambda: calls.__setitem__("n", calls["n"] + 1))

    class FakeResult:
        extractions = []

    monkeypatch.setenv("LANGEXTRACT_API_KEY", "fake-key")
    monkeypatch.setattr(deep_scan.lx, "extract", lambda **kwargs: FakeResult())
    entities, status = deep_scan.run_deep_scan("some text")
    assert status == "ok"
    assert calls["n"] == 1


def test_deep_scan_throttles_every_langextract_chunk_not_just_the_outer_call(monkeypatch):
    # lx.extract() can fan out into several concurrent per-chunk calls --
    # each chunk must throttle on its own, not just the outer call.
    from app import deep_scan

    calls = {"n": 0}
    monkeypatch.setattr(deep_scan.gemini_limiter, "acquire", lambda: calls.__setitem__("n", calls["n"] + 1))

    def fake_original(self, prompt, config):
        return f"handled: {prompt}"

    monkeypatch.setattr(deep_scan, "_original_process_single_prompt", fake_original)

    result_a = deep_scan.GeminiLanguageModel._process_single_prompt(object(), "chunk A", {})
    result_b = deep_scan.GeminiLanguageModel._process_single_prompt(object(), "chunk B", {})

    assert result_a == "handled: chunk A"
    assert result_b == "handled: chunk B"
    assert calls["n"] == 2


# ------------------------------------------------- shared retry (app/retry.py) ----
# time.sleep is faked so tests don't wait out the real backoff window.


def test_retry_backs_off_on_rate_limit_then_succeeds(monkeypatch):
    import httpx
    import openai

    from app import retry

    sleeps = []
    monkeypatch.setattr(retry.time, "sleep", lambda s: sleeps.append(s))

    calls = {"n": 0}
    fake_response = httpx.Response(
        status_code=429, request=httpx.Request("POST", "https://api.openai.com/v1/responses")
    )

    def flaky():
        calls["n"] += 1
        if calls["n"] < 3:
            raise openai.RateLimitError("rate limited", response=fake_response, body=None)
        return "recovered text"

    result = retry.call_with_backoff(flaky, label="test")
    assert result == "recovered text"
    assert calls["n"] == 3
    assert len(sleeps) == 2


def test_retry_backs_off_on_gemini_503_overload_then_succeeds(monkeypatch):
    # A Gemini 503 (google.genai.errors.ServerError) is transient too,
    # distinct exception type from the 429 case above.
    from google.genai import errors as genai_errors

    from app import retry

    sleeps = []
    monkeypatch.setattr(retry.time, "sleep", lambda s: sleeps.append(s))

    calls = {"n": 0}

    def flaky():
        calls["n"] += 1
        if calls["n"] < 2:
            raise genai_errors.ServerError(
                503,
                {"error": {"code": 503, "message": "high demand", "status": "UNAVAILABLE"}},
            )
        return "recovered text"

    result = retry.call_with_backoff(flaky, label="test")
    assert result == "recovered text"
    assert calls["n"] == 2
    assert sleeps == [3]


def test_retry_permanent_errors_fail_immediately_not_just_rate_limit(monkeypatch):
    # 401/404 are not transient -- must not retry.
    import httpx
    import openai
    from google.genai import errors as genai_errors

    from app import retry

    monkeypatch.setattr(retry.time, "sleep", lambda s: (_ for _ in ()).throw(
        AssertionError("must not sleep/retry on a permanent error")
    ))

    fake_401 = httpx.Response(
        status_code=401, request=httpx.Request("POST", "https://api.openai.com/v1/responses")
    )
    try:
        retry.call_with_backoff(
            lambda: (_ for _ in ()).throw(
                openai.AuthenticationError("bad key", response=fake_401, body=None)
            ),
            label="test",
        )
        assert False, "expected AuthenticationError to propagate"
    except openai.AuthenticationError:
        pass

    try:
        retry.call_with_backoff(
            lambda: (_ for _ in ()).throw(
                genai_errors.ClientError(404, {"error": {"code": 404, "message": "not found"}})
            ),
            label="test",
        )
        assert False, "expected ClientError to propagate"
    except genai_errors.ClientError:
        pass


def test_retry_gives_up_after_max_attempts(monkeypatch):
    import httpx
    import openai

    from app import retry

    monkeypatch.setattr(retry.time, "sleep", lambda s: None)
    fake_response = httpx.Response(
        status_code=429, request=httpx.Request("POST", "https://api.openai.com/v1/responses")
    )

    def always_limited():
        raise openai.RateLimitError("rate limited", response=fake_response, body=None)

    try:
        retry.call_with_backoff(always_limited, label="test")
        assert False, "expected RateLimitError to propagate after exhausting retries"
    except openai.RateLimitError:
        pass


def test_retry_non_transient_error_fails_immediately(monkeypatch):
    from app import retry

    calls = {"n": 0}

    def _boom():
        calls["n"] += 1
        raise ValueError("not a rate limit error")

    try:
        retry.call_with_backoff(_boom, label="test")
        assert False, "expected immediate failure, no retry"
    except ValueError:
        pass
    assert calls["n"] == 1


def test_ocr_image_via_api_rejects_unknown_engine():
    from app import ocr_api

    try:
        ocr_api.ocr_image_via_api("deepseek", b"fake")
        assert False, "expected ValueError for an unregistered engine"
    except ValueError:
        pass


def test_list_openai_style_models_filters_non_vision_families(monkeypatch):
    # /v1/models has no capability flag -- exclusion list must still filter
    # out whisper/embedding models from the OCR picker.
    from app import ocr_api

    class FakeModel:
        def __init__(self, id):
            self.id = id

    class FakeModels:
        def list(self):
            return [
                FakeModel("gpt-5.6"),
                FakeModel("whisper-1"),
                FakeModel("text-embedding-3-large"),
            ]

    class FakeClient:
        def __init__(self, **kwargs):
            self.models = FakeModels()

    monkeypatch.setenv("OPENAI_API_KEY", "fake-key")
    monkeypatch.setattr(ocr_api, "OpenAI", FakeClient)
    models, status = ocr_api.list_openai_style_models(api_key_env="OPENAI_API_KEY", base_url=None)
    assert status == "ok"
    assert models == ["gpt-5.6"]


def test_list_openai_style_models_reports_error_on_failure(monkeypatch):
    from app import ocr_api

    class FakeClient:
        def __init__(self, **kwargs):
            raise RuntimeError("boom")

    monkeypatch.setenv("OPENAI_API_KEY", "fake-key")
    monkeypatch.setattr(ocr_api, "OpenAI", FakeClient)
    models, status = ocr_api.list_openai_style_models(api_key_env="OPENAI_API_KEY", base_url=None)
    assert status == "skipped_error"
    assert models == []


# ------------------------------------------------------ OCR with boxes ----
# app/redact.py's scanned-PDF path -- ocr_words_via_api and its helpers.


def test_parse_bbox_json_strips_markdown_code_fence():
    from app.ocr_api import _parse_bbox_json

    raw = '```json\n[{"text": "hi", "box_2d": [1,2,3,4]}]\n```'
    assert _parse_bbox_json(raw) == [{"text": "hi", "box_2d": [1, 2, 3, 4]}]


def test_parse_bbox_json_plain_array_no_fence():
    from app.ocr_api import _parse_bbox_json

    assert _parse_bbox_json('[{"text": "hi", "box_2d": [1,2,3,4]}]') == [
        {"text": "hi", "box_2d": [1, 2, 3, 4]}
    ]


def test_parse_bbox_json_tolerates_raw_control_character_in_string():
    # Gemini sometimes emits a raw newline instead of "\n" -- must not
    # reject the whole response over it.
    from app.ocr_api import _parse_bbox_json

    raw = '[{"text": "line one\nline two", "box_2d": [1,2,3,4]}]'
    assert _parse_bbox_json(raw) == [
        {"text": "line one\nline two", "box_2d": [1, 2, 3, 4]}
    ]


def test_items_to_words_converts_gemini_normalized_coords():
    from app.ocr_api import _items_to_words

    items = [{"text": "Contact john@example.com", "box_2d": [100, 200, 150, 800]}]
    words = _items_to_words(items, width=1000, height=2000, coord_order="ymin_xmin_ymax_xmax", scale=1000)
    assert words == [("Contact john@example.com", (200.0, 200.0, 800.0, 300.0))]


def test_items_to_words_converts_openai_normalized_coords():
    from app.ocr_api import _items_to_words

    items = [{"text": "foo", "box_2d": [10, 20, 30, 40]}]
    words = _items_to_words(items, width=999, height=999, coord_order="xmin_ymin_xmax_ymax", scale=999)
    assert words == [("foo", (10.0, 20.0, 30.0, 40.0))]


def test_items_to_words_skips_malformed_entries():
    from app.ocr_api import _items_to_words

    items = [
        {"text": "", "box_2d": [1, 2, 3, 4]},  # empty text
        {"text": "ok", "box_2d": [1, 2, 3]},  # wrong length
        {"text": "ok", "box_2d": ["a", "b", "c", "d"]},  # non-numeric
        "not even a dict",
        {"text": "good", "box_2d": [0, 0, 10, 10]},
    ]
    words = _items_to_words(items, width=100, height=100, coord_order="xmin_ymin_xmax_ymax", scale=100)
    assert [w.text for w in words] == ["good"]


def test_items_to_words_tolerates_gemini_key_drift(monkeypatch):
    # Gemini sometimes emits "label"/"box" instead of "text"/"box_2d" --
    # every combination here must still parse.
    from app.ocr_api import _items_to_words

    items = [
        {"text": "line via text key", "box_2d": [1, 2, 3, 4]},
        {"label": "line via label key, box_2d key", "box_2d": [5, 6, 7, 8]},
        {"text": "line via text key, box key", "box": [9, 10, 11, 12]},
        {"label": "line via label key, box key", "box": [13, 14, 15, 16]},
    ]
    words = _items_to_words(items, width=100, height=100, coord_order="xmin_ymin_xmax_ymax", scale=100)
    assert [w.text for w in words] == [
        "line via text key",
        "line via label key, box_2d key",
        "line via text key, box key",
        "line via label key, box key",
    ]


def test_ocr_words_via_api_gemini_parses_real_sdk_call(monkeypatch):
    from app import ocr_api

    class FakeResponse:
        text = '[{"text": "Contact john@example.com", "box_2d": [100, 100, 200, 900]}]'

    class FakeModels:
        def generate_content(self, **kwargs):
            return FakeResponse()

    class FakeClient:
        def __init__(self, **kwargs):
            self.models = FakeModels()

    monkeypatch.setattr(ocr_api, "genai", type("G", (), {"Client": FakeClient})())
    monkeypatch.setattr(ocr_api, "_image_size", lambda png_bytes: (1000, 1000))
    monkeypatch.setenv("LANGEXTRACT_API_KEY", "fake-key")

    words = ocr_api.ocr_words_via_api("gemini", b"fake-png")
    assert len(words) == 1
    assert words[0].text == "Contact john@example.com"
    assert words[0].bbox == (100.0, 100.0, 900.0, 200.0)


def test_ocr_words_via_api_openai_parses_real_sdk_call(monkeypatch):
    from app import ocr_api

    class FakeResponse:
        output_text = '[{"text": "Contact john@example.com", "box_2d": [100, 100, 900, 200]}]'

    class FakeResponses:
        def create(self, **kwargs):
            return FakeResponse()

    class FakeClient:
        def __init__(self, **kwargs):
            self.responses = FakeResponses()

    monkeypatch.setattr(ocr_api, "OpenAI", FakeClient)
    monkeypatch.setattr(ocr_api, "_image_size", lambda png_bytes: (999, 999))
    monkeypatch.setenv("OPENAI_API_KEY", "fake-key")

    words = ocr_api.ocr_words_via_api("openai", b"fake-png")
    assert len(words) == 1
    assert words[0].text == "Contact john@example.com"
    assert words[0].bbox == (100.0, 100.0, 900.0, 200.0)


def test_ocr_words_via_api_no_key_raises_not_configured(monkeypatch):
    from app.ocr_api import OcrApiNotConfigured, ocr_words_via_api

    monkeypatch.delenv("XAI_API_KEY", raising=False)
    try:
        ocr_words_via_api("grok", b"fake-png")
        assert False, "expected OcrApiNotConfigured"
    except OcrApiNotConfigured as exc:
        assert exc.env_var == "XAI_API_KEY"


# ---------------------------------------------------- round 2 categories ----


def test_detects_pem_private_key(scan):
    resp = scan(
        "Server config:\n-----BEGIN RSA PRIVATE KEY-----\nMIIEpAIBAAKCAQEA...\n"
        "-----END RSA PRIVATE KEY-----\nDo not share."
    )
    assert "CRYPTO_PRIVATE_KEY" in entity_types(resp)


def test_detects_private_ip_cidr(scan):
    resp = scan(
        "Sơ đồ mạng: gateway nội bộ tại 10.20.5.1/24, subnet backup 192.168.1.0/24.",
        confidence_threshold=0.4,
    )
    assert "INFRA_NETWORK_MAP" in entity_types(resp)


def test_gps_coordinates_need_context_to_clear_default_threshold(scan):
    # Same shape appears in non-location text (e.g. pixel dims) without context.
    no_context = scan("Kích thước ảnh: 21.038300, 105.782900 pixel.", confidence_threshold=0.7)
    with_context = scan(
        "Toạ độ GPS của kho hàng: 21.038300, 105.782900.", confidence_threshold=0.3
    )
    assert "GPS_LOCATION" not in entity_types(no_context)
    assert "GPS_LOCATION" in entity_types(with_context)


def test_financial_credential_with_banking_context_outscores_bare(scan):
    # Longer gap here is deliberate — an earlier 30-char cap missed this phrasing.
    with_ctx = scan(
        "Mật khẩu ngân hàng của tài khoản công ty là: Xk9pL2.", confidence_threshold=0.3
    )
    without_ctx = scan("Mật khẩu wifi quán cafe là: Xk9pL2.", confidence_threshold=0.3)

    ctx_scores = [
        e["score"] for e in with_ctx.json()["detected_entities"] if e["entity_type"] == "FINANCIAL_CREDENTIAL"
    ]
    bare_scores = [
        e["score"]
        for e in without_ctx.json()["detected_entities"]
        if e["entity_type"] == "FINANCIAL_CREDENTIAL"
    ]
    assert ctx_scores, "banking-context password should be detected"
    assert bare_scores, "bare password assignment should still be detected (weakly)"
    assert max(ctx_scores) > max(bare_scores)


# ------------------------------------------------------------- deep scan ----
# Patches target app.scanning.run_deep_scan, the bound name it imports.


def test_deep_scan_off_by_default_never_calls_langextract(client, api_key, monkeypatch):
    def _boom(text, model_id=None):
        raise AssertionError("run_deep_scan must not run when deep_scan is unset")

    monkeypatch.setattr("app.scanning.run_deep_scan", _boom)
    resp = client.post("/api/v1/scan", json={"text": "hello"}, headers={"X-API-Key": api_key})
    assert resp.status_code == 200
    assert resp.json()["deep_scan_status"] is None


def test_deep_scan_merges_successful_extraction(client, api_key, monkeypatch):
    def _fake(text, model_id=None):
        entity = DetectedEntity(
            entity_type="HR_SENSITIVE_CONTENT",
            location=EntityLocation(start=0, end=5),
            text_val=text[0:5],
            score=0.6,
            context_snippet=text[0:5],
        )
        return [entity], "ok"

    monkeypatch.setattr("app.scanning.run_deep_scan", _fake)
    resp = client.post(
        "/api/v1/scan",
        json={"text": "Nhân viên bị kỷ luật.", "deep_scan": True},
        headers={"X-API-Key": api_key},
    )
    body = resp.json()
    assert body["deep_scan_status"] == "ok"
    assert "HR_SENSITIVE_CONTENT" in entity_types(resp)


def test_deep_scan_organization_examples_are_well_formed():
    # Guards against a future edit breaking the few-shot example data's shape.
    from app.deep_scan import EXAMPLES

    org_extractions = [
        ext
        for ex in EXAMPLES
        for ext in ex.extractions
        if ext.extraction_class == "ORGANIZATION"
    ]
    assert len(org_extractions) >= 2
    for ext in org_extractions:
        assert ext.extraction_text.startswith("Công ty")


# Every value class deep_scan's EXAMPLES should cover. Not derived from
# recognizers.yaml, so a future edit that drops a class's example fails loud.
_EXPECTED_DEEP_SCAN_VALUE_CLASSES = {
    "ORGANIZATION", "PERSON", "LOCATION",
    "EMAIL_ADDRESS", "PHONE_NUMBER", "URL", "IP_ADDRESS", "CREDIT_CARD",
    "IBAN_CODE", "CRYPTO", "MAC_ADDRESS", "US_SSN",
    "CONTRACT_ID", "INTERNAL_TAX_CODE", "FINANCIAL_METRIC", "EMPLOYEE_ID",
    "INFRA_SECRET", "IP_SENSITIVE_MARKER", "CRYPTO_PRIVATE_KEY",
    "INFRA_NETWORK_MAP", "GPS_LOCATION", "FINANCIAL_CREDENTIAL",
    "VN_NATIONAL_ID", "BANK_ACCOUNT_NUMBER", "FULL_ADDRESS",
}


def test_deep_scan_examples_are_grounded_and_cover_every_expected_class():
    # langextract silently drops an extraction whose text isn't an exact
    # substring of its example -- a typo here would fail silently.
    from app.deep_scan import EXAMPLES

    seen_classes = set()
    for ex in EXAMPLES:
        for ext in ex.extractions:
            assert ext.extraction_text in ex.text, (
                f"{ext.extraction_class!r} extraction_text is not a verbatim "
                f"substring of its example text -- would silently fail to "
                f"ground against the real API"
            )
            seen_classes.add(ext.extraction_class)

    missing = _EXPECTED_DEEP_SCAN_VALUE_CLASSES - seen_classes
    assert not missing, f"no example teaches these classes: {sorted(missing)}"


def test_deep_scan_overlap_types_matches_examples_coverage(monkeypatch):
    # _DEEP_SCAN_OVERLAP_TYPES must cover every class deep_scan can produce,
    # or an overlapping finding would sit duplicated instead of replacing it.
    from app import scanning

    assert _EXPECTED_DEEP_SCAN_VALUE_CLASSES <= scanning._DEEP_SCAN_OVERLAP_TYPES


def test_deep_scan_organization_merges_into_response(client, api_key, monkeypatch):
    def _fake(text, model_id=None):
        idx = text.index("Công ty")
        entity = DetectedEntity(
            entity_type="ORGANIZATION",
            location=EntityLocation(start=idx, end=idx + len("Công ty Cổ phần Đầu Tư Toàn Cầu")),
            text_val="Công ty Cổ phần Đầu Tư Toàn Cầu",
            score=0.6,
            context_snippet=text,
        )
        return [entity], "ok"

    monkeypatch.setattr("app.scanning.run_deep_scan", _fake)
    resp = client.post(
        "/api/v1/scan",
        json={
            "text": "đại diện Công ty Cổ phần Đầu Tư Toàn Cầu.",
            "deep_scan": True,
        },
        headers={"X-API-Key": api_key},
    )
    body = resp.json()
    assert body["deep_scan_status"] == "ok"
    entities = {(e["entity_type"], e["text_val"]) for e in body["detected_entities"]}
    assert ("ORGANIZATION", "Công ty Cổ phần Đầu Tư Toàn Cầu") in entities
    # Free-path truncated LOCATION must be replaced, not sit alongside it.
    assert ("LOCATION", "Toàn Cầu") not in entities


def test_deep_scan_overlap_dedup_does_not_drop_unrelated_nested_entities(
    client, api_key, monkeypatch
):
    # A PHONE_NUMBER nested inside an HR_SENSITIVE_CONTENT sentence is a
    # separate finding and must survive the overlap dedup above.
    text = "Nhân viên bị khiển trách do vi phạm, SĐT liên hệ 0912345678."

    def _fake(text, model_id=None):
        entity = DetectedEntity(
            entity_type="HR_SENSITIVE_CONTENT",
            location=EntityLocation(start=0, end=len(text)),
            text_val=text,
            score=0.6,
            context_snippet=text,
        )
        return [entity], "ok"

    monkeypatch.setattr("app.scanning.run_deep_scan", _fake)
    resp = client.post(
        "/api/v1/scan",
        json={"text": text, "deep_scan": True, "confidence_threshold": 0.3},
        headers={"X-API-Key": api_key},
    )
    entities = {(e["entity_type"], e["text_val"]) for e in resp.json()["detected_entities"]}
    assert ("HR_SENSITIVE_CONTENT", text) in entities
    assert ("PHONE_NUMBER", "0912345678") in entities


def test_anonymize_masks_deep_scan_organization(client, api_key, monkeypatch):
    text = "đại diện Công ty Cổ phần Đầu Tư Toàn Cầu."

    def _fake(text, model_id=None):
        idx = text.index("Công ty")
        entity = DetectedEntity(
            entity_type="ORGANIZATION",
            location=EntityLocation(start=idx, end=idx + len("Công ty Cổ phần Đầu Tư Toàn Cầu")),
            text_val="Công ty Cổ phần Đầu Tư Toàn Cầu",
            score=0.6,
            context_snippet=text,
        )
        return [entity], "ok"

    monkeypatch.setattr("app.scanning.run_deep_scan", _fake)
    resp = client.post(
        "/api/v1/scan",
        json={"text": text, "deep_scan": True, "anonymize": True, "confidence_threshold": 0.3},
        headers={"X-API-Key": api_key},
    )
    anonymized = resp.json()["anonymized_content"]["text"]
    assert "Công ty Cổ phần Đầu Tư Toàn Cầu" not in anonymized
    assert "<ORGANIZATION>" in anonymized


def test_anonymize_does_not_swallow_sentence_around_hr_sensitive_content(
    client, api_key, monkeypatch
):
    # HR_SENSITIVE_CONTENT flags a topic, not a value -- must not feed into
    # AnonymizerEngine, or it'd swallow the nested PHONE_NUMBER into one tag.
    text = "Nhân viên bị khiển trách do vi phạm, SĐT liên hệ 0912345678."

    def _fake(text, model_id=None):
        entity = DetectedEntity(
            entity_type="HR_SENSITIVE_CONTENT",
            location=EntityLocation(start=0, end=len(text)),
            text_val=text,
            score=0.6,
            context_snippet=text,
        )
        return [entity], "ok"

    monkeypatch.setattr("app.scanning.run_deep_scan", _fake)
    resp = client.post(
        "/api/v1/scan",
        json={"text": text, "deep_scan": True, "anonymize": True, "confidence_threshold": 0.3},
        headers={"X-API-Key": api_key},
    )
    anonymized = resp.json()["anonymized_content"]["text"]
    assert anonymized == "Nhân viên bị khiển trách do vi phạm, SĐT liên hệ <PHONE_NUMBER>."


def test_deep_scan_failure_falls_back_to_regex_results(client, api_key, monkeypatch):
    monkeypatch.setattr("app.scanning.run_deep_scan", lambda text, model_id=None: ([], "skipped_error"))
    resp = client.post(
        "/api/v1/scan",
        json={"text": "Contact test@example.com", "deep_scan": True},
        headers={"X-API-Key": api_key},
    )
    body = resp.json()
    assert resp.status_code == 200
    assert body["deep_scan_status"] == "skipped_error"
    assert "EMAIL_ADDRESS" in entity_types(resp)


def test_deep_scan_quota_exceeded_after_cap(client, monkeypatch):
    from app.pages import MAX_DEEP_SCAN_PER_KEY

    monkeypatch.setattr("app.scanning.run_deep_scan", lambda text, model_id=None: ([], "ok"))

    email = f"quota-{uuid.uuid4().hex[:8]}@sensen.dev"
    key = client.post("/register", json={"email": email}).json()["api_key"]

    for _ in range(MAX_DEEP_SCAN_PER_KEY):
        resp = client.post(
            "/api/v1/scan", json={"text": "hello", "deep_scan": True}, headers={"X-API-Key": key}
        )
        assert resp.json()["deep_scan_status"] == "ok"

    resp = client.post(
        "/api/v1/scan", json={"text": "hello", "deep_scan": True}, headers={"X-API-Key": key}
    )
    assert resp.json()["deep_scan_status"] == "skipped_quota_exceeded"


def test_run_deep_scan_retries_once_on_langextract_race_then_succeeds(monkeypatch):
    # langextract's Gemini provider intermittently raises on schema
    # validation, then succeeds on immediate retry.
    from app import deep_scan

    calls = {"n": 0}

    def flaky_extract(**kwargs):
        calls["n"] += 1
        if calls["n"] == 1:
            raise RuntimeError("simulated langextract schema-validation race")

        class FakeResult:
            extractions = []

        return FakeResult()

    monkeypatch.setenv("LANGEXTRACT_API_KEY", "fake-key")
    monkeypatch.setattr(deep_scan.lx, "extract", flaky_extract)
    entities, status = deep_scan.run_deep_scan("some text")
    assert status == "ok"
    assert calls["n"] == 2


def test_run_deep_scan_gives_up_after_max_attempts(monkeypatch):
    from app import deep_scan

    calls = {"n": 0}

    def always_fails(**kwargs):
        calls["n"] += 1
        raise RuntimeError("simulated permanent langextract failure")

    monkeypatch.setenv("LANGEXTRACT_API_KEY", "fake-key")
    monkeypatch.setattr(deep_scan.lx, "extract", always_fails)
    entities, status = deep_scan.run_deep_scan("some text")
    assert status == "skipped_error"
    assert calls["n"] == deep_scan._MAX_ATTEMPTS


def test_run_deep_scan_backs_off_on_rate_limit_then_succeeds(monkeypatch):
    # A real RPM limit (InferenceRuntimeError wrapping ClientError(429) in
    # .original) needs a backoff sleep, unlike the schema race above.
    from google.genai import errors as genai_errors

    from app import deep_scan

    sleeps = []
    monkeypatch.setattr(deep_scan.time, "sleep", lambda s: sleeps.append(s))

    calls = {"n": 0}

    def flaky_extract(**kwargs):
        calls["n"] += 1
        if calls["n"] == 1:
            raise deep_scan.lx.exceptions.InferenceRuntimeError(
                "Gemini API error: 429",
                original=genai_errors.ClientError(
                    429, {"error": {"code": 429, "message": "rate limited"}}
                ),
            )

        class FakeResult:
            extractions = []

        return FakeResult()

    monkeypatch.setenv("LANGEXTRACT_API_KEY", "fake-key")
    monkeypatch.setattr(deep_scan.lx, "extract", flaky_extract)
    entities, status = deep_scan.run_deep_scan("some text")
    assert status == "ok"
    assert calls["n"] == 2
    assert sleeps == [deep_scan.BACKOFF_SECONDS[0]]


def test_deep_scan_model_override_passed_through(client, api_key, monkeypatch):
    captured = {}

    def _fake(text, model_id=None):
        captured["model_id"] = model_id
        return [], "ok"

    monkeypatch.setattr("app.scanning.run_deep_scan", _fake)
    client.post(
        "/api/v1/scan",
        json={"text": "hello", "deep_scan": True, "model": "gemini-3.7-flash"},
        headers={"X-API-Key": api_key},
    )
    assert captured["model_id"] == "gemini-3.7-flash"


def test_deep_scan_rejects_unusable_model_override(monkeypatch):
    # Validated before any network call — an image/tts/etc. model id under
    # the same "gemini-" prefix is rejected without spending an API call.
    from app.deep_scan import run_deep_scan

    monkeypatch.setenv("LANGEXTRACT_API_KEY", "fake-key-for-validation-test")
    entities, status = run_deep_scan("some text", model_id="gemini-3-pro-image")
    assert status == "skipped_error"
    assert entities == []


def test_deep_scan_models_endpoint_reports_no_key(client, api_key, monkeypatch):
    monkeypatch.delenv("LANGEXTRACT_API_KEY", raising=False)
    resp = client.get("/api/v1/deep_scan/models", headers={"X-API-Key": api_key})
    assert resp.status_code == 200
    body = resp.json()
    assert body["status"] == "skipped_no_key"
    assert body["models"] == []
    assert body["default_model"]


def test_deep_scan_models_endpoint_requires_auth(client):
    resp = client.get("/api/v1/deep_scan/models")
    assert resp.status_code == 422  # missing required X-API-Key header


# ------------------------------------------ real-document findings fix ----


def test_vn_mobile_number_detected_as_phone(scan):
    resp = scan("Điện thoại: 0932843132, liên hệ ngay.")
    assert "PHONE_NUMBER" in entity_types(resp)


def test_vn_national_id_context_outscores_bare_digits(scan):
    with_ctx = scan("Số CCCD: 051195344431 cấp tại Hà Nội.", confidence_threshold=0.3)
    without_ctx = scan("Mã đơn hàng: 051195344431 đã xử lý.", confidence_threshold=0.3)

    ctx_scores = [
        e["score"] for e in with_ctx.json()["detected_entities"] if e["entity_type"] == "VN_NATIONAL_ID"
    ]
    bare_scores = [
        e["score"]
        for e in without_ctx.json()["detected_entities"]
        if e["entity_type"] == "VN_NATIONAL_ID"
    ]
    assert ctx_scores, "CCCD-shaped number with context should be detected"
    assert not bare_scores or max(ctx_scores) > max(bare_scores)


def test_exact_span_duplicate_across_categories_keeps_only_highest_score(scan):
    # PhoneRecognizer also matches a VN national ID's digit shape at a
    # lower score -- the duplicate on the exact same span must be dropped.
    resp = scan("Số CCCD: 051195344431 cấp tại Hà Nội.", confidence_threshold=0.3)
    entities = resp.json()["detected_entities"]
    same_span = [e for e in entities if e["location"]["start"] == 9 and e["location"]["end"] == 21]
    assert {e["entity_type"] for e in same_span} == {"VN_NATIONAL_ID"}


# --------------------------------------- 4 categories added while fixing ----


def test_detects_old_cmnd_9_digit_with_context(scan):
    resp = scan("CMND/Hộ chiếu số: 371775053 do CA Kiên Giang cấp.", confidence_threshold=0.3)
    assert "VN_NATIONAL_ID" in entity_types(resp)


def test_bare_9_digit_number_needs_context_to_clear_default_threshold(scan):
    resp = scan("Mã đơn hàng: 123456789 đã xử lý.", confidence_threshold=0.7)
    assert "VN_NATIONAL_ID" not in entity_types(resp)


def test_detects_contract_id_number_first_format(scan):
    # "14/2014/HĐCBL" -- serial/year/abbreviation, reverse of HD-YYYY-XXX.
    resp = scan("Hợp đồng cấp bảo lãnh số 14/2014/HĐCBL đã ký.", confidence_threshold=0.5)
    assert "CONTRACT_ID" in entity_types(resp)


def test_detects_bank_account_number_with_context(scan):
    resp = scan("Số tài khoản: 0091000099412 tại Vietcombank.", confidence_threshold=0.3)
    assert "BANK_ACCOUNT_NUMBER" in entity_types(resp)


def test_detects_grouped_bank_account_number(scan):
    resp = scan("Tài khoản số: 341.0100.00041 dùng để nhận tiền bảo lãnh.", confidence_threshold=0.3)
    assert "BANK_ACCOUNT_NUMBER" in entity_types(resp)


def test_detects_full_street_address(scan):
    resp = scan(
        "Địa chỉ: Số 89, đường 3/2, phường Vĩnh Bảo, Thành phố Rạch Giá, tỉnh Kiên Giang.",
        confidence_threshold=0.3,
    )
    entities = resp.json()["detected_entities"]
    addr = next((e for e in entities if e["entity_type"] == "FULL_ADDRESS"), None)
    assert addr is not None, "expected a FULL_ADDRESS hit"
    assert "Kiên Giang" in addr["text_val"] and "Số 89" in addr["text_val"]


def test_full_address_still_matches_when_admin_levels_are_abbreviated(scan):
    # Real addresses often skip/abbreviate levels ("TP" instead of "thành phố").
    resp = scan(
        "Địa chỉ: Số 09 Huỳnh Tịnh Của, Phường Vĩnh Thanh Vân, TP Rạch Giá, Kiên Giang.",
        confidence_threshold=0.3,
    )
    assert "FULL_ADDRESS" in entity_types(resp)


def test_vietnamese_ner_now_correct_instead_of_noisy(scan):
    resp = scan(
        "Một bên là Ông/Bà: Nguyễn Xuân Hùng, Quốc tịch: Việt Nam. Chức vụ: Giám đốc điều hành.",
        confidence_threshold=0.3,
    )
    entities = {(e["entity_type"], e["text_val"]) for e in resp.json()["detected_entities"]}
    assert ("PERSON", "Nguyễn Xuân Hùng") in entities
    assert ("LOCATION", "Việt Nam") in entities
    # Fragments like "Ông/Bà" or "Chức vụ" must not also get tagged.
    garbage = {(t, v) for t, v in entities if v not in ("Nguyễn Xuân Hùng", "Việt Nam")}
    assert not garbage, f"unexpected NER hits: {garbage}"


def test_vietnamese_ner_score_is_graduated_not_flat(scan):
    # Score must vary meaningfully, not a flat 0.6 for every kept result.
    resp = scan(
        "Một bên là Ông/Bà: Nguyễn Xuân Hùng, Quốc tịch: Việt Nam.",
        confidence_threshold=0.1,
    )
    scores = {
        e["entity_type"]: e["score"]
        for e in resp.json()["detected_entities"]
        if e["entity_type"] in ("PERSON", "LOCATION")
    }
    assert scores["PERSON"] > 0.6
    assert scores["LOCATION"] > 0.6


def test_vietnamese_ner_bullet_initial_common_word_filtered_at_normal_threshold(scan):
    # "Được"/"Xét" are capitalized only from starting a bulleted clause.
    resp = scan(
        "- Được trả lương vào ngày 05 hàng tháng.\n- Xét nâng lương định kỳ 12 tháng/lần.",
        confidence_threshold=0.5,
    )
    entities = {(e["entity_type"], e["text_val"]) for e in resp.json()["detected_entities"]}
    noise = {(t, v) for t, v in entities if t in ("PERSON", "LOCATION", "ORGANIZATION")}
    assert not noise, f"unexpected NER hits: {noise}"


def test_name_split_across_newline_by_underthesea_still_recovered(scan):
    # underthesea can fuse a name with the next line's label across a "\n" --
    # must still resolve back to the source text.
    resp = scan(
        "Nhân viên: Trần Thị Hoa\nSố CCCD: 038196045678\nPhòng ban: Kế toán",
        confidence_threshold=0.3,
    )
    entities = {(e["entity_type"], e["text_val"]) for e in resp.json()["detected_entities"]}
    assert ("PERSON", "Trần Thị Hoa") in entities


def test_bare_newline_before_single_word_still_penalized_as_sentence_initial(scan):
    # A bare "\n" (no preceding punctuation) must still count as a sentence
    # boundary for the capitalization penalty.
    resp = scan(
        "Server config:\nInternal API key: sk-live-51Hh8x9AbCdEfGhIjKlMnOpQrStUvWxYz01234567",
        confidence_threshold=0.5,
    )
    entities = {(e["entity_type"], e["text_val"]) for e in resp.json()["detected_entities"]}
    noise = {(t, v) for t, v in entities if t in ("PERSON", "LOCATION", "ORGANIZATION")}
    assert not noise, f"unexpected NER hits: {noise}"


def test_find_token_does_not_desync_on_a_distant_unrelated_occurrence():
    # An unbounded search could latch onto a distant unrelated occurrence of
    # a repeated token, permanently desyncing every later lookup -- the
    # search must be bounded to _MAX_TOKEN_LOOKAHEAD and simply fail instead.
    from app.vi_ner import _MAX_TOKEN_LOOKAHEAD, _find_token

    filler = "x" * (_MAX_TOKEN_LOOKAHEAD + 50)
    text = f"{filler} Kiên Giang là tỉnh đẹp."

    found = _find_token(text, "Kiên Giang", 0)
    assert found is None, (
        "must not silently match the distant unrelated occurrence — that's "
        "exactly the jump that caused the real desync"
    )


def test_collapse_repeated_punctuation_shrinks_long_runs():
    # A long dot-leader run (VN form blank-fill lines) makes underthesea
    # silently skip tokenizing real text right after it -- collapse first.
    from app.vi_ner import _REPEATED_PUNCT_KEEP, _collapse_repeated_punctuation

    long_run = "." * 60
    text = f"Số tiền: {long_run}\nHọ tên: Nguyễn Văn A."
    collapsed, mapping = _collapse_repeated_punctuation(text)

    assert "." * _REPEATED_PUNCT_KEEP in collapsed
    assert "." * (_REPEATED_PUNCT_KEEP + 1) not in collapsed
    assert "Nguyễn Văn A" in collapsed
    assert len(collapsed) == len(mapping)


def test_collapse_repeated_punctuation_leaves_short_runs_alone():
    from app.vi_ner import _collapse_repeated_punctuation

    text = "Giá: 1,234.56 -- xem thêm... rồi kết luận."
    collapsed, mapping = _collapse_repeated_punctuation(text)
    assert collapsed == text
    assert mapping == list(range(len(text)))


def test_collapse_repeated_punctuation_mapping_resolves_to_real_offsets():
    from app.vi_ner import _collapse_repeated_punctuation

    text = "A" + "-" * 50 + "Nguyễn Thị Mĩnh" + "-" * 50 + "B"
    collapsed, mapping = _collapse_repeated_punctuation(text)
    idx = collapsed.find("Nguyễn Thị Mĩnh")
    assert idx != -1
    start = mapping[idx]
    end = mapping[idx + len("Nguyễn Thị Mĩnh") - 1] + 1
    assert text[start:end] == "Nguyễn Thị Mĩnh"


def test_compose_mappings_chains_two_transformations_with_dash_one_propagation():
    from app.vi_ner import _compose_mappings

    # first_stage: some earlier transform's text -> original text
    first_stage = [0, 1, -1, 2, 3]  # position 2 was inserted (no origin)
    # second_stage: final text -> first_stage's text
    second_stage = [4, 3, 2, 0]
    composed = _compose_mappings(second_stage, first_stage)
    assert composed == [3, 2, -1, 0]


def test_fused_word_no_longer_falsely_tagged(scan):
    # Some PDF exports drop the space glyph between word pairs -- a fused
    # run like "Chếđộlàm" must not look like one Title Case name.
    resp = scan("Điều 2: Chếđộlàm việc theo quy định công ty.", confidence_threshold=0.3)
    entities = {(e["entity_type"], e["text_val"]) for e in resp.json()["detected_entities"]}
    noise = {(t, v) for t, v in entities if t in ("PERSON", "LOCATION", "ORGANIZATION")}
    assert not noise, f"unexpected NER hits on a fused-word run: {noise}"


def test_fused_name_recovered_via_syllable_segmentation(scan):
    resp = scan(
        "Và một bên là Ông/Bà: Trịnh SỹThành Quốc tịch: Đài Loan.",
        confidence_threshold=0.3,
    )
    entities = {(e["entity_type"], e["text_val"]) for e in resp.json()["detected_entities"]}
    assert ("PERSON", "Trịnh SỹThành") in entities


def test_vietnamese_ner_no_hallucination_on_fragment_only_text(scan):
    resp = scan(
        "Mọi thắc mắc vui lòng liên hệ trong giờ hành chính để được hỗ trợ.",
        confidence_threshold=0.3,
    )
    assert not (entity_types(resp) & {"PERSON", "ORGANIZATION", "LOCATION"})


def test_vietnamese_day_names_rejected_not_tagged_location(scan):
    # Day names ("Thứ Bảy", "Chủ Nhật") are a closed set, never PII.
    resp = scan(
        "Nghỉ hàng tuần 02 ngày (Thứ Bảy, Chủ Nhật); nghỉ phép năm 12 ngày.",
        confidence_threshold=0.1,
    )
    entities = {(e["entity_type"], e["text_val"]) for e in resp.json()["detected_entities"]}
    noise = {(t, v) for t, v in entities if t in ("PERSON", "LOCATION", "ORGANIZATION")}
    assert not noise, f"day names should never be tagged as entities: {noise}"


def test_admin_unit_prefix_corrects_type_to_location(scan):
    # underthesea gets the span right but sometimes tags it PERSON --
    # "Phường" is a reliable signal to correct the type to LOCATION.
    resp = scan(
        "Địa chỉ: 93A/51/223 Ngô Gia Tự, Phường Thủ Thiêm, TP. Hồ Chí Minh.",
        confidence_threshold=0.3,
    )
    entities = {(e["entity_type"], e["text_val"]) for e in resp.json()["detected_entities"]}
    assert ("LOCATION", "Phường Thủ Thiêm") in entities
    assert ("PERSON", "Phường Thủ Thiêm") not in entities


def test_tax_code_context_outscores_bare_digits(scan):
    with_ctx = scan("Mã số thuế doanh nghiệp: 1234567890", confidence_threshold=0.2)
    without_ctx = scan("Random id: 1234567890", confidence_threshold=0.2)

    ctx_scores = [
        e["score"] for e in with_ctx.json()["detected_entities"] if e["entity_type"] == "INTERNAL_TAX_CODE"
    ]
    bare_scores = [
        e["score"]
        for e in without_ctx.json()["detected_entities"]
        if e["entity_type"] == "INTERNAL_TAX_CODE"
    ]
    assert ctx_scores, "context-boosted tax code should be detected at threshold 0.2"
    assert not bare_scores or max(ctx_scores) > max(bare_scores)

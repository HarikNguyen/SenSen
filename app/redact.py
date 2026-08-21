"""Real redaction for PDF/DOCX/TXT -- content genuinely removed/blacked
out, not just masked text like app/scanning.py's anonymize=true.

PDF: digital-text pages use pymupdf's search_for + apply_redactions (real
deletion); scanned pages are OCR'd then blacked out and the page image
replaced. DOCX: paragraph text is spliced out of the runs (real XML
deletion); embedded images share the scanned-PDF OCR-and-blackout path.
Not covered: DOCX comments, tracked-changes, text boxes, OLE objects,
document properties. TXT: same masking as anonymize=true, as a file.

Safety fallback: an entity that can't be located fails the whole request
(RedactionFailed) rather than shipping a silently under-redacted file.
"""

import io
import logging
from typing import Optional

import pymupdf
import pytesseract
from docx import Document
from PIL import Image, ImageDraw
from presidio_analyzer import AnalyzerEngine, RecognizerResult
from presidio_anonymizer import AnonymizerEngine

from app.deep_scan import run_deep_scan
from app.extract import UnsupportedFileType
from app.ocr_api import OcrApiError, OcrApiNotConfigured, OcrWord, ocr_words_via_api
from app.schemas import DetectedEntity, EntityLocation
from app.scanning import (
    CONTEXT_WINDOW,
    _ANONYMIZE_EXCLUDED_TYPES,
    _drop_lower_scored_exact_duplicates,
    _drop_regex_ner_entities_overlapped_by_deep_scan,
)

logger = logging.getLogger("sensen.redact")

# Redaction on a scanned page is at least as heavy as OCR alone.
MAX_REDACT_PAGES = 20
OCR_LANGUAGES = "vie+eng"  # matches app/extract.py's local-OCR language set

# Formats this server can decode for redaction; anything else fails loudly.
_CONTENT_TYPE_TO_PIL_FORMAT = {
    "image/png": "PNG",
    "image/jpeg": "JPEG",
    "image/gif": "GIF",
    "image/bmp": "BMP",
    "image/tiff": "TIFF",
}


class RedactionFailed(ValueError):
    """A page's sensitive content couldn't be located for redaction."""


def redact_pdf(
    raw: bytes,
    *,
    analyzer: AnalyzerEngine,
    confidence_threshold: float,
    ocr_engine: str = "local",
    ocr_model: Optional[str] = None,
    deep_scan: bool = False,
    deep_scan_model: Optional[str] = None,
) -> bytes:
    """Raises UnsupportedFileType (bad file/OCR config) or RedactionFailed
    (entity not locatable)."""
    try:
        doc = pymupdf.open(stream=raw, filetype="pdf")
    except pymupdf.FileDataError as exc:
        raise UnsupportedFileType(
            f"This file isn't a valid PDF (couldn't be opened: {exc}). "
            f"Check it isn't corrupted or empty."
        ) from exc

    try:
        if len(doc) > MAX_REDACT_PAGES:
            raise UnsupportedFileType(
                f"This PDF has {len(doc)} pages, over the {MAX_REDACT_PAGES}-page "
                f"redaction limit -- redaction is at least as CPU/API-heavy as OCR."
            )
        for page_index in range(len(doc)):
            page = doc[page_index]
            if page.get_text().strip():
                _redact_digital_page(page, analyzer, confidence_threshold, deep_scan, deep_scan_model)
            else:
                _redact_scanned_page(
                    doc,
                    page_index,
                    analyzer,
                    confidence_threshold,
                    ocr_engine,
                    ocr_model,
                    deep_scan,
                    deep_scan_model,
                )
        # without these, a redacted image comes out uncompressed and bloated
        return doc.tobytes(garbage=4, deflate=True, clean=True)
    finally:
        doc.close()


def redact_docx(
    raw: bytes,
    *,
    analyzer: AnalyzerEngine,
    confidence_threshold: float,
    ocr_engine: str = "local",
    ocr_model: Optional[str] = None,
    deep_scan: bool = False,
    deep_scan_model: Optional[str] = None,
) -> bytes:
    """Text via run-splicing, embedded images via OCR-and-blackout."""
    try:
        document = Document(io.BytesIO(raw))
    except Exception as exc:
        # python-docx doesn't guarantee one exception type for a malformed file.
        raise UnsupportedFileType(
            f"This file isn't a valid .docx (couldn't be opened: {exc}). "
            f"Check it isn't corrupted."
        ) from exc

    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for section in document.sections:
        paragraphs.extend(section.header.paragraphs)
        paragraphs.extend(section.footer.paragraphs)

    for paragraph in paragraphs:
        _redact_docx_paragraph(
            paragraph, analyzer, confidence_threshold, deep_scan, deep_scan_model
        )

    _redact_docx_images(
        document,
        analyzer,
        confidence_threshold,
        ocr_engine,
        ocr_model,
        deep_scan,
        deep_scan_model,
    )

    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def redact_txt(
    raw: bytes,
    *,
    analyzer: AnalyzerEngine,
    anonymizer: AnonymizerEngine,
    confidence_threshold: float,
    deep_scan: bool = False,
    deep_scan_model: Optional[str] = None,
) -> bytes:
    """Same masking as anonymize=true, returned as a file instead of JSON."""
    text = raw.decode("utf-8", errors="replace")
    entities = _detect_entities_for_redaction(
        text, analyzer, confidence_threshold, deep_scan, deep_scan_model
    )
    anonymize_results = [
        RecognizerResult(
            entity_type=e.entity_type, start=e.location.start, end=e.location.end, score=e.score
        )
        for e in entities
    ]
    anon = anonymizer.anonymize(text=text, analyzer_results=anonymize_results)
    return anon.text.encode("utf-8")


def _detect_entities_for_redaction(
    text: str,
    analyzer: AnalyzerEngine,
    confidence_threshold: float,
    deep_scan: bool,
    deep_scan_model: Optional[str],
) -> list[DetectedEntity]:
    """Same detection/dedup/deep-scan logic as run_scan()'s anonymize path,
    plus _merge_overlapping_entities since this module splices text
    directly instead of going through AnonymizerEngine."""
    results = analyzer.analyze(text=text, language="en", score_threshold=confidence_threshold)
    results = _drop_lower_scored_exact_duplicates(results)
    entities = [
        DetectedEntity(
            entity_type=r.entity_type,
            location=EntityLocation(start=r.start, end=r.end),
            text_val=text[r.start : r.end],
            score=round(r.score, 3),
            context_snippet=text[
                max(0, r.start - CONTEXT_WINDOW) : min(len(text), r.end + CONTEXT_WINDOW)
            ],
        )
        for r in results
    ]

    if deep_scan:
        deep_entities, _status = run_deep_scan(text, model_id=deep_scan_model)
        entities = _drop_regex_ner_entities_overlapped_by_deep_scan(entities, deep_entities)
        entities.extend(deep_entities)

    entities = [e for e in entities if e.entity_type not in _ANONYMIZE_EXCLUDED_TYPES]
    return _merge_overlapping_entities(entities, text)


def _merge_overlapping_entities(entities: list[DetectedEntity], text: str) -> list[DetectedEntity]:
    """Union overlapping spans (e.g. nested URL inside EMAIL_ADDRESS) so
    splicing both doesn't corrupt the text or under-redact either side."""
    if not entities:
        return entities
    ordered = sorted(entities, key=lambda e: e.location.start)
    merged: list[DetectedEntity] = [ordered[0]]
    for entity in ordered[1:]:
        last = merged[-1]
        if entity.location.start >= last.location.end:
            merged.append(entity)
            continue
        start = last.location.start
        end = max(last.location.end, entity.location.end)
        winner = last if last.score >= entity.score else entity
        merged[-1] = DetectedEntity(
            entity_type=winner.entity_type,
            location=EntityLocation(start=start, end=end),
            text_val=text[start:end],
            score=winner.score,
            context_snippet=text[max(0, start - CONTEXT_WINDOW) : min(len(text), end + CONTEXT_WINDOW)],
        )
    return merged


def _redact_digital_page(
    page,
    analyzer: AnalyzerEngine,
    confidence_threshold: float,
    deep_scan: bool,
    deep_scan_model: Optional[str],
) -> None:
    text = page.get_text()
    entities = _detect_entities_for_redaction(text, analyzer, confidence_threshold, deep_scan, deep_scan_model)
    if not entities:
        return

    for entity in entities:
        rects = page.search_for(entity.text_val)
        if not rects:
            raise RedactionFailed(
                f"Page {page.number + 1}: found sensitive text "
                f"'{entity.text_val[:30]}...' during analysis but couldn't "
                f"locate it on the page for redaction -- refusing to produce "
                f"a possibly under-redacted PDF."
            )
        for rect in rects:
            page.add_redact_annot(rect, fill=(0, 0, 0))
    page.apply_redactions()


def _redact_scanned_page(
    doc,
    page_index: int,
    analyzer: AnalyzerEngine,
    confidence_threshold: float,
    ocr_engine: str,
    ocr_model: Optional[str],
    deep_scan: bool,
    deep_scan_model: Optional[str],
) -> None:
    page = doc[page_index]
    pix = page.get_pixmap(dpi=200)  # matches app/extract.py's _ocr_page
    mode = "RGBA" if pix.alpha else "RGB"
    img = Image.frombytes(mode, (pix.width, pix.height), pix.samples)

    img = _ocr_and_redact_image(
        img,
        label=f"Page {page_index + 1}",
        analyzer=analyzer,
        confidence_threshold=confidence_threshold,
        ocr_engine=ocr_engine,
        ocr_model=ocr_model,
        deep_scan=deep_scan,
        deep_scan_model=deep_scan_model,
        not_found_error=UnsupportedFileType,
    )
    if img is None:
        return  # genuinely blank page, or nothing sensitive found -- untouched

    redacted_buf = io.BytesIO()
    img.save(redacted_buf, format="PNG")

    rect = page.rect
    doc.delete_page(page_index)
    new_page = doc.new_page(pno=page_index, width=rect.width, height=rect.height)
    new_page.insert_image(new_page.rect, stream=redacted_buf.getvalue())


def _redact_docx_paragraph(
    paragraph,
    analyzer: AnalyzerEngine,
    confidence_threshold: float,
    deep_scan: bool,
    deep_scan_model: Optional[str],
) -> None:
    text = paragraph.text
    if not text.strip():
        return
    entities = _detect_entities_for_redaction(
        text, analyzer, confidence_threshold, deep_scan, deep_scan_model
    )
    if not entities:
        return

    run_spans = _run_spans(paragraph)
    # reverse order so an earlier span's offsets never shift mid-splice
    for entity in sorted(entities, key=lambda e: e.location.start, reverse=True):
        _splice_entity_from_runs(entity.location.start, entity.location.end, run_spans)


def _run_spans(paragraph) -> list[tuple[int, int, object]]:
    """[(start, end, run), ...] offsets into paragraph.text."""
    spans = []
    cursor = 0
    for run in paragraph.runs:
        start = cursor
        cursor += len(run.text)
        spans.append((start, cursor, run))
    return spans


def _splice_entity_from_runs(
    entity_start: int, entity_end: int, run_spans: list[tuple[int, int, object]]
) -> None:
    """Deletes text[entity_start:entity_end] from whichever run(s) it overlaps."""
    for run_start, run_end, run in run_spans:
        if run_start >= entity_end or run_end <= entity_start:
            continue
        local_start = max(entity_start, run_start) - run_start
        local_end = min(entity_end, run_end) - run_start
        run.text = run.text[:local_start] + run.text[local_end:]


def _redact_docx_images(
    document,
    analyzer: AnalyzerEngine,
    confidence_threshold: float,
    ocr_engine: str,
    ocr_model: Optional[str],
    deep_scan: bool,
    deep_scan_model: Optional[str],
) -> None:
    image_parts = [
        part
        for part in document.part.related_parts.values()
        if part.content_type.startswith("image/")
    ]
    for i, part in enumerate(image_parts, start=1):
        label = f"Embedded image {i}"
        pil_format = _CONTENT_TYPE_TO_PIL_FORMAT.get(part.content_type)
        if pil_format is None:
            raise UnsupportedFileType(
                f"{label} is a {part.content_type} image, which this server "
                f"can't inspect for redaction -- refusing to produce a "
                f"possibly under-redacted file. Supported image types: "
                f"{', '.join(sorted(_CONTENT_TYPE_TO_PIL_FORMAT))}."
            )
        try:
            img = Image.open(io.BytesIO(part.blob))
            img.load()
        except Exception as exc:
            raise UnsupportedFileType(f"{label} couldn't be decoded: {exc}") from exc

        redacted = _ocr_and_redact_image(
            img,
            label=label,
            analyzer=analyzer,
            confidence_threshold=confidence_threshold,
            ocr_engine=ocr_engine,
            ocr_model=ocr_model,
            deep_scan=deep_scan,
            deep_scan_model=deep_scan_model,
            not_found_error=UnsupportedFileType,
        )
        if redacted is None:
            continue  # nothing sensitive found in this image -- untouched

        buf = io.BytesIO()
        if pil_format == "JPEG" and redacted.mode in ("RGBA", "P"):  # JPEG has no alpha
            redacted = redacted.convert("RGB")
        redacted.save(buf, format=pil_format)
        part._blob = buf.getvalue()


def _ocr_and_redact_image(
    img: "Image.Image",
    *,
    label: str,
    analyzer: AnalyzerEngine,
    confidence_threshold: float,
    ocr_engine: str,
    ocr_model: Optional[str],
    deep_scan: bool,
    deep_scan_model: Optional[str],
    not_found_error: type,
) -> Optional["Image.Image"]:
    """OCR + blackout, or None if nothing to redact. Shared by
    `_redact_scanned_page` and `_redact_docx_images`."""
    if ocr_engine == "local":
        try:
            data = pytesseract.image_to_data(
                img, lang=OCR_LANGUAGES, output_type=pytesseract.Output.DICT
            )
        except pytesseract.TesseractNotFoundError:
            raise not_found_error(
                f"{label} needs OCR to redact, but the server's tesseract "
                f"binary isn't installed. See README."
            ) from None
        except pytesseract.TesseractError as exc:
            raise not_found_error(
                f"Tesseract OCR failed on {label.lower()} (often missing "
                f"language data, e.g. tesseract-ocr-vie): {exc}"
            ) from None
        text, spans = _tesseract_data_to_text_and_spans(data)
    else:
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        try:
            lines = ocr_words_via_api(ocr_engine, buf.getvalue(), ocr_model)
        except OcrApiNotConfigured as exc:
            raise not_found_error(
                f"{label} needs OCR to redact via the '{ocr_engine}' API, but "
                f"the server has no {exc.env_var} configured. See README."
            ) from None
        except OcrApiError as exc:
            raise not_found_error(
                f"OCR-with-boxes call failed on {label.lower()} ('{ocr_engine}'): {exc}"
            ) from None
        if not lines:
            raise RedactionFailed(
                f"{label}: '{ocr_engine}' returned no usable text/boxes -- "
                f"refusing to produce a possibly under-redacted file. Try "
                f"ocr_engine=local or a different engine."
            )
        text, spans = _lines_to_text_and_spans(lines)

    if not text.strip():
        return None  # genuinely blank image, nothing to redact

    entities = _detect_entities_for_redaction(text, analyzer, confidence_threshold, deep_scan, deep_scan_model)
    if not entities:
        return None

    boxes = []
    for entity in entities:
        words = _overlapping_words(entity.location.start, entity.location.end, spans)
        if not words:
            raise RedactionFailed(
                f"{label}: found sensitive text '{entity.text_val[:30]}...' "
                f"during OCR but couldn't locate it on the image for "
                f"redaction -- refusing to produce a possibly under-redacted file."
            )
        boxes.append(_union_bbox(words))

    redacted = img.copy()
    draw = ImageDraw.Draw(redacted)
    for box in boxes:
        draw.rectangle(box, fill=(0, 0, 0))
    return redacted


def _tesseract_data_to_text_and_spans(
    data: dict,
) -> tuple[str, list[tuple[int, int, OcrWord]]]:
    """Reconstructs text + each word's [start, end) span for interval matching."""
    text_parts: list[str] = []
    spans: list[tuple[int, int, OcrWord]] = []
    cursor = 0
    prev_line_key = None
    for i in range(len(data["text"])):
        word_text = data["text"][i].strip()
        if not word_text:
            continue
        line_key = (data["block_num"][i], data["par_num"][i], data["line_num"][i])
        if prev_line_key is not None:
            sep = "\n" if line_key != prev_line_key else " "
            text_parts.append(sep)
            cursor += len(sep)
        prev_line_key = line_key

        left, top = data["left"][i], data["top"][i]
        width, height = data["width"][i], data["height"][i]
        word = OcrWord(text=word_text, bbox=(left, top, left + width, top + height))
        start = cursor
        text_parts.append(word_text)
        cursor += len(word_text)
        spans.append((start, cursor, word))

    return "".join(text_parts), spans


def _lines_to_text_and_spans(lines: list[OcrWord]) -> tuple[str, list[tuple[int, int, OcrWord]]]:
    """Same idea as _tesseract_data_to_text_and_spans, for cloud engines' per-line boxes."""
    text_parts: list[str] = []
    spans: list[tuple[int, int, OcrWord]] = []
    cursor = 0
    for i, line in enumerate(lines):
        if i > 0:
            text_parts.append("\n")
            cursor += 1
        start = cursor
        text_parts.append(line.text)
        cursor += len(line.text)
        spans.append((start, cursor, line))
    return "".join(text_parts), spans


def _overlapping_words(
    entity_start: int, entity_end: int, spans: list[tuple[int, int, OcrWord]]
) -> list[OcrWord]:
    return [word for (s, e, word) in spans if s < entity_end and e > entity_start]


def _union_bbox(words: list[OcrWord]) -> tuple[float, float, float, float]:
    x0 = min(w.bbox[0] for w in words)
    y0 = min(w.bbox[1] for w in words)
    x1 = max(w.bbox[2] for w in words)
    y1 = max(w.bbox[3] for w in words)
    return (x0, y0, x1, y1)

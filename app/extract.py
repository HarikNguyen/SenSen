"""PDF/DOCX/TXT -> text. Digital text layer first, OCR fallback (local
Tesseract by default; cloud via `ocr_engine` is opt-in, see app/ocr_api.py).
"""

import io
import logging
from pathlib import Path
from typing import Optional

import pymupdf
import pytesseract
from docx import Document
from PIL import Image, ImageFilter

from app.ocr_api import OcrApiError, OcrApiNotConfigured, ocr_image_via_api

logger = logging.getLogger("sensen.extract")

# Caps OCR cost (CPU for local, paid calls for cloud engines) on a many-page scan.
MAX_OCR_PAGES = 20
OCR_LANGUAGES = "vie+eng"


class UnsupportedFileType(ValueError):
    pass


def extract_text(
    filename: str, raw: bytes, ocr_engine: str = "local", ocr_model: Optional[str] = None
) -> tuple[str, str, int, str]:
    """Returns (text, file_type, total_pages, processing_mode). `ocr_model`
    is independent from deep_scan's own `model` param."""
    suffix = Path(filename or "").suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(raw, ocr_engine, ocr_model)
    if suffix == ".docx":
        return _extract_docx(raw)
    if suffix == ".txt":
        return raw.decode("utf-8", errors="replace"), "text", 1, "direct_text_extraction"

    raise UnsupportedFileType(
        f"Unsupported file type '{suffix or '(none)'}'. MVP supports .pdf "
        f"(digital text or scanned via OCR), .docx and .txt."
    )


def _extract_pdf(
    raw: bytes, ocr_engine: str = "local", ocr_model: Optional[str] = None
) -> tuple[str, str, int, str]:
    try:
        doc = pymupdf.open(stream=raw, filetype="pdf")
    except pymupdf.FileDataError as exc:
        # also covers an empty upload (EmptyFileError subclasses this)
        raise UnsupportedFileType(
            f"This file isn't a valid PDF (couldn't be opened: {exc}). "
            f"Check it isn't corrupted or empty."
        ) from exc
    try:
        pages = [page.get_text() for page in doc]
        # per-page, not document-wide: a PDF can mix digital-text and scanned pages
        if all(p.strip() for p in pages):
            return "\n".join(pages), "pdf", len(pages), "direct_text_extraction"

        if len(doc) > MAX_OCR_PAGES:
            raise UnsupportedFileType(
                f"This PDF has no extractable text layer (likely scanned) and "
                f"has {len(doc)} pages, over the {MAX_OCR_PAGES}-page OCR limit "
                f"— OCR is CPU-heavy, capped to keep it bounded on modest hardware."
            )

        # Only OCR pages that lack a text layer; reuse the rest.
        merged = [p if p.strip() else _ocr_page(page, ocr_engine, ocr_model) for p, page in zip(pages, doc)]
        merged_text = "\n".join(merged)
        if not merged_text.strip():
            raise UnsupportedFileType(
                "This PDF has no extractable text layer, and OCR found no "
                "readable text either (likely blank pages or very low scan quality)."
            )
        return merged_text, "pdf", len(pages), f"ocr_{ocr_engine}"
    finally:
        doc.close()


def _ocr_page(page, ocr_engine: str = "local", ocr_model: Optional[str] = None) -> str:
    pix = page.get_pixmap(dpi=200)  # enough for OCR accuracy without a full-res render
    mode = "RGBA" if pix.alpha else "RGB"
    img = Image.frombytes(mode, (pix.width, pix.height), pix.samples)

    if ocr_engine != "local":
        # Cloud models handle noisy/blurry scans natively; skip the local filter.
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        try:
            return ocr_image_via_api(ocr_engine, buf.getvalue(), ocr_model)
        except OcrApiNotConfigured as exc:
            logger.error("OCR API requested (%s) but not configured: %s", ocr_engine, exc)
            raise UnsupportedFileType(
                f"This PDF has no extractable text layer and needs OCR via "
                f"the '{ocr_engine}' API, but the server has no {exc.env_var} "
                f"configured. See README."
            ) from None
        except OcrApiError as exc:
            logger.error("OCR API call failed (%s): %s", ocr_engine, exc)
            raise UnsupportedFileType(
                f"This PDF has no extractable text layer, and the "
                f"'{ocr_engine}' OCR API call failed: {exc}"
            ) from None

    # size=7: smaller kernels leave photocopy speckle that fragments OCR output
    img = img.filter(ImageFilter.MedianFilter(size=7))
    try:
        return pytesseract.image_to_string(img, lang=OCR_LANGUAGES)
    except pytesseract.TesseractNotFoundError:
        logger.error("tesseract binary isn't installed -- see README")
        raise UnsupportedFileType(
            "This PDF has no extractable text layer and needs OCR, but the "
            "server's tesseract binary isn't installed. See README."
        ) from None
    except pytesseract.TesseractError as exc:
        # usually missing language data (e.g. tesseract-ocr-vie)
        logger.error("tesseract call failed: %s", exc)
        raise UnsupportedFileType(
            "This PDF has no extractable text layer and needs OCR, but the "
            "server's tesseract call failed (often missing language data, "
            "e.g. tesseract-ocr-vie). See README."
        ) from None


def _extract_docx(raw: bytes) -> tuple[str, str, int, str]:
    try:
        document = Document(io.BytesIO(raw))
    except Exception as exc:
        # python-docx raises different exception types for malformed input
        raise UnsupportedFileType(
            f"This file isn't a valid .docx (couldn't be opened: {exc}). "
            f"Check it isn't corrupted."
        ) from exc
    parts = [p.text for p in document.paragraphs]

    # Flatten tables to pipe-delimited rows so cell values keep context for scoring.
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))

    return "\n".join(parts), "docx", 1, "direct_text_extraction"
